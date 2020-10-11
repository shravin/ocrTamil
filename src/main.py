import os
import traceback

from src.combine import combineDocsFromDir, combineFormattedDocsFromDir
from docx import Document

from src.constants import PNG_FILE_EXTENSION, DOCX_FILE_EXTENSION, ANNOTATED_FILE_PREFIX, FORMATTED_FILE_PREFIX
from src.doc_text import render_doc_text, get_document_paragraphs, write_annotated_image, FeatureType
from src.utils import is_file_dir_present


def write_formatted_document(full_text_annotation, header_name, formatted_doc_file_name):
    print("Writing the formatted document with file name: {}".format(formatted_doc_file_name))
    paragraphs, lines = get_document_paragraphs(full_text_annotation)
    document = Document()
    document.add_heading(header_name, 0)

    for para in paragraphs:
        doc_page = document.add_paragraph()
        doc_page.add_run(para)

    document.save(formatted_doc_file_name)


def main(image_dir, doc_dir, ref_dir, combined_filename):
    # the below line gets only the files in the image_dir
    file_list = [file for file in os.listdir(image_dir)
             if os.path.isfile(os.path.join(image_dir, file))]
    counter = 0
    total_files = len(file_list)
    for file_name in file_list:
        try:
            counter += 1
            input_file_with_path="{}/{}".format(image_dir, file_name)
            print("Processing file no: {} of {} with filename: {}".format(counter, total_files, input_file_with_path))
            file_name_without_extension = os.path.splitext(file_name)[0]
            formatted_doc_file_name = "{}/{}{}{}".format(doc_dir, FORMATTED_FILE_PREFIX, file_name_without_extension,
                                                         DOCX_FILE_EXTENSION)
            doc_file_name = "{}/{}{}".format(doc_dir, file_name_without_extension, DOCX_FILE_EXTENSION)
            annotated_file_name = "{}/{}{}{}".format(ref_dir, ANNOTATED_FILE_PREFIX, file_name_without_extension,
                                                     PNG_FILE_EXTENSION)

            if is_file_dir_present(doc_file_name):
                print("Skipping processing file: {} as the file {} is present.".format(file_name, doc_file_name))
                continue

            texts, full_text_annotation, image = render_doc_text(input_file_with_path)
            ocr_text = texts[0].description

            write_doc_without_formatting(doc_file_name, file_name_without_extension, ocr_text)
            write_formatted_document(full_text_annotation, file_name_without_extension, formatted_doc_file_name)
            write_annotated_image(image, full_text_annotation, FeatureType.WORD, ref_dir, file_name_without_extension)
            write_annotated_image(image, full_text_annotation, FeatureType.PARA, ref_dir, file_name_without_extension)

        except Exception as e:
            print("Error in processing for filename: {}".format(file_name))
            traceback.print_exc()
            print(e)
            pass

    if combined_filename is None:
        combined_filename = "combinedDoc.docx"

    combineDocsFromDir(doc_dir, combined_filename)
    combineFormattedDocsFromDir(doc_dir, "{}{}".format(FORMATTED_FILE_PREFIX, combined_filename))


def write_doc_without_formatting(doc_file_name, file_name_without_extension, ocr_text):
    document = Document()
    document.add_heading(file_name_without_extension, 0)
    doc_page = document.add_paragraph()
    doc_page.add_run(ocr_text)
    document.add_page_break()
    document.save(doc_file_name)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description='OCR for tamil novels - Provide dir')
    parser.add_argument('--input_imagedir', metavar='path', required=True,
                        help='the path to input image dirs for OCR')
    # parser.add_argument('--output_docdir', metavar='path', required=True,
    #                     help='the path to output doc dir after OCR')
    # parser.add_argument('--reference_imgdir', metavar='path', required=True,
    #                     help='the path to output reference image dir after OCR')
    parser.add_argument('--combined_filename', metavar='path', required=False,
                        help='the filename of the single file name', default=None)

    args = parser.parse_args()

    image_dir = args.input_imagedir
    # doc_dir = args.output_docdir
    # reference_imgdir = args.reference_imgdir
    combined_filename = args.combined_filename

    image_dir_present = is_file_dir_present(image_dir)
    doc_dir = image_dir + "/outdir"
    reference_imgdir = image_dir + "/refdir"
    os.mkdir(doc_dir)
    os.mkdir(reference_imgdir)

    doc_dir_present = is_file_dir_present(doc_dir)
    reference_imgdir_present = is_file_dir_present(reference_imgdir)

    if doc_dir_present and image_dir_present and reference_imgdir_present:
        main(image_dir=image_dir, doc_dir=doc_dir, ref_dir=reference_imgdir, combined_filename=combined_filename)
    else:
        print("Program exits without doing anything as requisite dirs are not present. The program expects "
              "the dirs to be present")

