import os
from docx import Document
from natsort import natsorted

from src.utils import is_file_dir_present


def combineDocsFromDir(inputDirPath, combined_filename="combinedDoc.docx"):
    print("Combining documents into a single file: ", combined_filename)
    file_list = os.listdir(inputDirPath)
    target_document = Document()
    sorted_filelist = natsorted(file_list)
    for file_name in sorted_filelist:
        print("Combining file: {}".format(file_name))
        source_document = Document("{}/{}".format(inputDirPath, file_name))
        for paragraph in source_document.paragraphs:
            text = paragraph.text
            target_document.add_paragraph(text)
        target_document.add_page_break()
    target_document.save("{}/{}".format(inputDirPath, combined_filename))


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description='OCR for tamil novels - Provide dir')
    parser.add_argument('--inputDirPath', metavar='path', required=True,
                        help='the path to input doc dir to combine')
    parser.add_argument('--combined_filename', metavar='path', required=True,
                        help='combined_filename')

    args = parser.parse_args()

    inputDirPath = args.inputDirPath
    combined_filename = args.combined_filename

    inputDirPath_present = is_file_dir_present(inputDirPath)

    if inputDirPath_present:
        combineDocsFromDir(inputDirPath=inputDirPath, combined_filename=combined_filename)
    else:
        print("Program exits without doing anything as requisite dirs are not present. The program expects "
              "the dirs to be present")
